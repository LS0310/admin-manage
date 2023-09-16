
import requests
from bs4 import BeautifulSoup
from fake_useragent import UserAgent
import os
import docx
import urllib.request
import re
 
def main():
    # 使用fake_useragent随机生成User-Agent伪造请求头
    ua = UserAgent()
    headers = {'User-Agent': ua.random}
 
    url = 'https://www.ynufe.edu.cn/pub/yjsyxz/zs/index.htm'
    res = requests.get(url, headers=headers)
 
    # 解析HTML文件
    soup = BeautifulSoup(res.text, 'html.parser', from_encoding='utf-8')
 
 
    for i in range(4):
        if i == 0:
            """
            解析招生动态部分
            """
            news_items = soup.select('.ui-index-item.col-md-6.col-sm-12.col-xs-12')[i].select('.news-item')
            # 创建文件夹
            if not os.path.exists("结果/招生动态"):
                os.makedirs("结果/招生动态")
            os.chdir("结果/招生动态")
        elif i == 1:
            """
            解析博士招生部分
            """
            news_items = soup.select('.ui-index-item.col-md-6.col-sm-12.col-xs-12')[i].select('.news-item')
            # 创建文件夹
            if not os.path.exists("../博士招生"):
                os.makedirs("../博士招生")
            os.chdir("../博士招生")
        elif i == 2:
            """
            解析硕士招生部分
            """
            news_items = soup.select('.ui-index-item.col-md-6.col-sm-12.col-xs-12')[i].select('.news-item')
            # 创建文件夹
            if not os.path.exists("../硕士招生"):
                os.makedirs("../硕士招生")
            os.chdir("../硕士招生")
        elif i == 3:
            """
            解析港澳台招生部分
            """
            news_items = soup.select('.ui-index-item.col-md-6.col-sm-12.col-xs-12')[i].select('.news-item')
            # 创建文件夹
            if not os.path.exists("../港澳台招生"):
                os.makedirs("../港澳台招生")
            os.chdir("../港澳台招生")
 
        for item in news_items:
            # 获取链接
            parts = url.split("/")
            del parts[-1]   # 删除最后一部分index.htm
            new_url = "/".join(parts)  # 组成https://www.ynufe.edu.cn/pub/yjsyxz/zs
 
            link = new_url+ "/" +item.select_one('a')['href']
 
            if link.endswith("pdf"):    #对pdf结尾单独处理
                item_text = str(item).encode('ISO-8859-1').decode('utf-8')
                item_text = re.sub(r'<.*?>', '', item_text)  # 去除所有标签
                item_text = re.sub(r'\s+', ' ', item_text)  # 替换连续的空白字符为一个空格
                match = re.search(r'\s*(\w+)\.{3}', item_text)
                if match:
                    result = match.group(1)
                    filename = result+".pdf"
                else:
                    filename = "文件名提取失败.pdf"
                urllib.request.urlretrieve(link, filename)    # 下载文件
            else:
                # 访问链接并获取html内容
                res1 = requests.get(link, headers=headers)
                res1.encoding = 'utf-8'  # 设置网页编码为utf-8
                html_content = res1.content.decode('utf-8') # 对获取的HTML内容进行解码
                soup1 = BeautifulSoup(html_content, 'html.parser') # 对HTML内容进行解析
 
                parts = link.split("/")
                del parts[-1]   # 删除最后一部分index.htm
                new_link = "/".join(parts)  
                # 此时soup1为招生动态的每一个子页面内容
                save_word(soup1,new_link)
 
 
 
"""
保存方式：
    遍历正文部分进行处理，并且拥有参数"当前文件夹"
        内容中遇到文本即存入,遇到表格则传递给表格函数存入,遇到链接则判断是htm还是pdf,htm则递归至"当前文件夹+标题"子文件夹
"""
def save_word(soup:BeautifulSoup,url:str):
    """
    将传入的页面保存在当前目录文件夹中,通过os.chdir指定
    :param soup对象
    :param url目前所在文档链接
    """
 
    title = soup.select_one('.news-cont-info-article .title').text.strip()
    # 创建word和写入标题
    doc = docx.Document()
    doc.add_heading(title, 0)
 
    # 处理class="page-news-cont"的主内容
    page_news_cont = soup.find(class_='page-news-cont')
    # 遍历page_news_cont标签中的所有子标签
    for tag in page_news_cont.descendants:
        # 如果是<p>标签则调用text_save函数
        if tag.name == 'p':
            text_save(tag,title,doc,url)
        # 如果是<table>标签则调用table_save函数
        elif tag.name == 'table':
            table_save(tag,title,doc,url)
        else:
            pass
 
    doc.save(title+'.docx')
    print("保存成功:"+title)
 
 
 
def text_save(tag,title,doc,url):
    # 处理<p>标签的代码,对应网页中一段话，直接保存，遇到链接做额外处理
    text = ""
    for child in tag.children:
        if child.name == 'span':
            text += child.text
        elif child.name == 'a':
            text += child.text
            link = child.get('href')
            a_save(None,title,doc,url,link,child.text)
 
    doc.add_paragraph(text)
 
 
def table_save(tag,title,doc,url):
    # 处理<table>标签的代码
    data = []
    # 遍历table标签中的所有tr标签
    for tr in tag.find_all('tr'):
        row = []
        # 遍历当前tr标签中的所有td标签
        for td in tr.find_all('td'):
            # 如果当前td标签中有<a>标签，则将<a>标签中的文本添加到row数组中，然后对<a>标签进行额外处理
            if td.find('a') is not None:
                text = td.find('a').text.strip()
                row.append(text)
                link = td.find('a')['href'].strip()
                a_save(tag, title, doc, url, link, text)
            # 如果当前td标签中有<span>标签，则将<span>标签中的文本添加到row数组中
            elif td.find('span') is not None:
                row.append(td.find('span').text.strip())
            # 否则，将当前td标签中的文本添加到row数组中
            else:
                row.append(td.text.strip())
        # 将当前行数据添加到data数组中
        data.append(row)
 
     
    # 创建一个新的表格对象(网站中有时候有空表格，这时候跳过)
    try:
        table = doc.add_table(rows=len(data), cols=max(len(row) for row in data))
    except:
        return
 
    # 遍历data数组中的每一行数据，并将每个单元格的值填入表格中
    for i in range(len(data)):
        for j in range(len(data[i])):
            table.cell(i, j).text = data[i][j]
 
 
 
def a_save(tag,title,doc,from_url,link,text):
    """
    处理<a>标签的代码
    :param tag:从save_word传递来的tag标签，暂时没用上
    :param title:目前所在文件的文件名
    :param doc:目前打开的word文档对象
    :param from_url:目前所在文档的网页链接
    :param link:目前要处理的a标签href链接,用于判断直接访问还是和from_url连接
    :param text:目前要处理的a标签的文本
    """
 
    # 将目前位置临时改为文件名，如"结果/招生动态/xxx名单/"
    if not os.path.exists(title):
        os.makedirs(title)
    os.chdir(title)
 
 
    if link.startswith("http"):
        # a标签的href以http开头，直接访问，不需要拼接前置
        link=link
    else:
        # a标签的href是相对路径，需要拼接前置
        link=from_url+'/'+link
         
    if link.endswith("pdf"):
        urllib.request.urlretrieve(link, text+".pdf")    # 下载文件
 
    # 返回之前的目录，继续其他文件
    os.chdir("..")
 
main()